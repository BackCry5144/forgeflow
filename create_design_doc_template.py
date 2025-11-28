# create_design_doc_template.py
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_template():
    doc = Document()
    
    # 기본 폰트 설정 (맑은 고딕)
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # [타이틀]
    title = doc.add_heading(level=0)
    run = title.add_run('[ {{SCREEN_NAME}} ] 화면 설계서')
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    # [1. 개요]
    doc.add_heading('1. 개요 및 목적', level=1)
    table_info = doc.add_table(rows=1, cols=2)
    table_info.style = 'Table Grid'
    table_info.rows[0].cells[0].text = "컴포넌트 ID"
    table_info.rows[0].cells[1].text = "{{COMPONENT_NAME}}"
    
    doc.add_paragraph('')
    doc.add_paragraph('{{DESCRIPTION}}')
    doc.add_paragraph('')

    # [2. 화면 예시]
    doc.add_heading('2. 화면 예시', level=1)
    p = doc.add_paragraph('{{SCREENSHOT}}') # 이미지 들어갈 자리
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    # [3. UI 구조] (핵심 변경 포인트 🔥)
    doc.add_heading('3. UI 구조', level=1)
    # 코드가 이 태그를 찾아서 '구조도 표'로 교체합니다.
    doc.add_paragraph('{{UI_STRUCTURE}}') 
    doc.add_paragraph('')

    # [4. 상태 관리 명세]
    doc.add_heading('4. 상태 관리 명세 (State)', level=1)
    table_state = doc.add_table(rows=1, cols=4)
    table_state.style = 'Table Grid'
    
    # 🔥 [태그 추가] 코드가 이 표를 찾을 수 있게 숨겨진 태그 추가
    # 첫 번째 셀에 태그를 넣고, 나중에 코드가 이를 지웁니다.
    headers = ["{{TABLE:STATE}} 변수명", "데이터 타입", "초기값", "설명"]
    
    for i, h in enumerate(headers):
        cell = table_state.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    # [5. 이벤트 명세]
    doc.add_heading('5. 이벤트 및 로직 명세', level=1)
    table_event = doc.add_table(rows=1, cols=3)
    table_event.style = 'Table Grid'
    
    # 🔥 [태그 추가]
    headers_evt = ["{{TABLE:EVENT}} UI 요소", "이벤트", "로직 상세 설명"]
    
    for i, h in enumerate(headers_evt):
        cell = table_event.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    # [6. 동작 흐름] (핵심 변경 포인트 🔥)
    doc.add_heading('6. 화면 동작 순서 (User Flow)', level=1)
    # 코드가 이 태그를 찾아서 'Normal Flow 표'로 교체합니다.
    doc.add_paragraph('{{USER_FLOW}}')

    # 저장
    output_path = 'backend/templates/design_spec_template_V2.docx'
    # 경로 없으면 생성
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    
    doc.save(output_path)
    print(f"✅ 템플릿 생성 완료: {output_path}")

if __name__ == "__main__":
    try:
        create_template()
    except ImportError:
        print("❌ python-docx 가 설치되지 않았습니다. 'pip install python-docx'를 실행하세요.")