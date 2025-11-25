# create_template.py

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_design_template():
    # 1. 문서 생성
    doc = Document()
    
    # ---------------------------------------------------------
    # 스타일 설정 (기본 폰트: 맑은 고딕)
    # ---------------------------------------------------------
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # ---------------------------------------------------------
    # 타이틀: [ {{SCREEN_NAME}} ] 화면 설계서
    # ---------------------------------------------------------
    title = doc.add_heading(level=0)
    run = title.add_run('[ {{SCREEN_NAME}} ] 화면 설계서')
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('') # 공백

    # ---------------------------------------------------------
    # 1. 기본 정보 (개요)
    # ---------------------------------------------------------
    h1 = doc.add_heading('1. 개요 및 목적', level=1)
    
    # 표 생성 (1행 2열) - 컴포넌트 ID 표시용
    table_info = doc.add_table(rows=1, cols=2)
    table_info.style = 'Table Grid'
    
    row = table_info.rows[0]
    row.cells[0].text = "컴포넌트 ID"
    row.cells[0].paragraphs[0].runs[0].font.bold = True
    row.cells[1].text = "{{COMPONENT_NAME}}"
    
    doc.add_paragraph('') # 공백
    
    p_desc = doc.add_paragraph()
    run_desc = p_desc.add_run('{{DESCRIPTION}}')
    
    doc.add_paragraph('')

    # ---------------------------------------------------------
    # 2. 화면 예시 (스크린샷)
    # ---------------------------------------------------------
    doc.add_heading('2. 화면 예시', level=1)
    p_shot = doc.add_paragraph('{{SCREENSHOT}}')
    p_shot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')

    # ---------------------------------------------------------
    # 3. UI 구조
    # ---------------------------------------------------------
    doc.add_heading('3. UI 구조', level=1)
    doc.add_paragraph('{{UI_STRUCTURE}}')
    
    doc.add_paragraph('')

    # ---------------------------------------------------------
    # 4. 상태 관리 (State Specification) - tables[0]
    # ---------------------------------------------------------
    doc.add_heading('4. 상태 관리 명세 (State)', level=1)
    
    # 빈 표 생성 (헤더만 존재)
    # DocumentService가 tables[0]을 찾아서 여기에 데이터를 채웁니다.
    table_state = doc.add_table(rows=1, cols=4)
    table_state.style = 'Table Grid'
    
    # 헤더 설정
    headers_state = ["변수명", "데이터 타입", "초기값", "설명"]
    hdr_cells = table_state.rows[0].cells
    for i, text in enumerate(headers_state):
        hdr_cells[i].text = text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_paragraph('')

    # ---------------------------------------------------------
    # 5. 이벤트 핸들러 (Event Handlers) - tables[1]
    # ---------------------------------------------------------
    doc.add_heading('5. 이벤트 및 로직 명세', level=1)
    
    # 빈 표 생성 (헤더만 존재)
    # DocumentService가 tables[1]을 찾아서 여기에 데이터를 채웁니다.
    table_event = doc.add_table(rows=1, cols=3)
    table_event.style = 'Table Grid'
    
    # 헤더 설정
    headers_event = ["함수명", "트리거", "로직 상세 설명"]
    hdr_cells = table_event.rows[0].cells
    for i, text in enumerate(headers_event):
        hdr_cells[i].text = text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------------------------------------------------
    # 파일 저장
    # ---------------------------------------------------------
    # templates 폴더가 없으면 생성
    if not os.path.exists('templates'):
        os.makedirs('templates')
        
    save_path = os.path.join('templates', 'design_spec_template.docx')
    doc.save(save_path)
    print(f"✅ 템플릿 생성 완료: {save_path}")

if __name__ == "__main__":
    try:
        create_design_template()
    except Exception as e:
        print(f"❌ 오류 발생: {e}")
        print("pip install python-docx 명령어로 라이브러리를 설치했는지 확인해주세요.")