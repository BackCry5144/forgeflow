# backend/services/document_service.py

import json
import os
import logging
import asyncio
import traceback
from io import BytesIO
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# GeminiClient 사용 (SSL 우회 지원)
from services.gemini_client import get_gemini_client, GeminiClientError

from utils.doc_prompts import get_design_spec_prompt, get_test_plan_prompt, get_user_manual_prompt

logger = logging.getLogger(__name__)


class DocumentService:
    def __init__(self):
        # GeminiClient 사용 (SSL 우회 지원)
        self.client = get_gemini_client()
        
        # 설정
        self.max_quota_retries = 3
        self.retry_delay_seconds = 30

        logger.info(f"DocumentService initialized with GeminiClient: {self.client.model_name}")

    async def _call_with_retry(self, prompt: str, operation_name: str = "LLM Call"):
        """LLM 호출 + 재시도 로직 (GeminiClient 사용)"""
        try:
            logger.info(f"🤖 [{operation_name}] Calling Gemini API...")
            
            # GeminiClient의 generate_content_async 사용
            text = await self.client.generate_content_async(
                prompt=prompt,
                temperature=0.2,
                max_output_tokens=8192,
                timeout=120
            )
            
            logger.info(f"📨 [{operation_name}] Response received ({len(text)} chars)")
            return text
            
        except GeminiClientError as e:
            logger.error(f"❌ [{operation_name}] GeminiClient Error: {e.error_type}: {e.message}")
            raise Exception(f"{e.error_type}: {e.message}")
        except Exception as e:
            logger.error(f"❌ [{operation_name}] Failed: {type(e).__name__}: {e}")
            raise e

    async def generate_design_doc(
        self, 
        screen_name: str, 
        react_code: str, 
        wizard_data: dict, 
        images: List[dict] = None
    ) -> BytesIO:
        """개발자용 화면 설계서 생성 (Word)"""
        logger.info(f"📄 Generating design spec for: {screen_name}")
        
        # 1. LLM 데이터 추출
        try:
            prompt = get_design_spec_prompt(react_code, wizard_data)
            logger.info(f"📝 Prompt generated: {len(prompt)} chars")
            
            text = await self._call_with_retry(prompt, "Design Spec")
            
            # JSON 파싱
            text = text.strip()
            if text.startswith("```json"): text = text[7:]
            if text.startswith("```"): text = text[3:]
            if text.endswith("```"): text = text[:-3]
            design_data = json.loads(text)
            logger.info("✅ LLM Data Extraction Success")
        except Exception as e:
            logger.error(f"❌ LLM Extraction Failed: {type(e).__name__}: {e}")
            traceback.print_exc()
            design_data = {"basic_info": {"screen_name": screen_name, "description": "분석 실패"}}

        # 2. Word 생성 (wizard_data 추가 전달)
        return self._create_design_docx(design_data, wizard_data, images)

    def _create_design_docx(self, data: dict, wizard_data: dict, images: List[dict] = None) -> BytesIO:
        base_dir = os.path.dirname(os.path.dirname(__file__))
        template_path = os.path.join(base_dir, 'templates', 'design_spec_template.docx')
        
        doc = Document(template_path) if os.path.exists(template_path) else Document()

        # (A) 텍스트 치환
        info = data.get('basic_info', {})
        replacements = {
            "{{SCREEN_NAME}}": str(info.get('screen_name', data.get('screen_name', ''))),
            "{{COMPONENT_NAME}}": str(info.get('component_name', '')),
            "{{DESCRIPTION}}": str(info.get('description', '')),
        }

        for paragraph in doc.paragraphs:
            for key, val in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(val))

        # (B) 이미지 삽입 ({{SCREENSHOT}})
        for paragraph in doc.paragraphs:
            if "{{SCREENSHOT}}" in paragraph.text:
                paragraph.text = "" 
                if images:
                    for img in images:
                        try:
                            run = paragraph.add_run()
                            run.add_picture(BytesIO(img['bytes']), width=Inches(6.0))
                            # 이미지 라벨 추가
                            run.add_text(f"\n[{img['label']}]\n")
                        except: pass
                else:
                    paragraph.text = "(스크린샷 없음)"
                break

        # (C) UI 구조 도식화 (Tree Grid 스타일) 🔥
        # wizard_data에서 컴포넌트 상세 목록을 가져옴
        raw_components = wizard_data.get('step3', {}).get('components', [])
        self._generate_ui_structure_table(doc, data.get('layout_structure', []), raw_components)

        # (D) Normal Flow (동작 순서)
        self._generate_user_flow_table(doc, data.get('user_flow', []))

        # (E) 고정 테이블 채우기 (태그 기반 검색)
        state_table = self._find_table_by_tag(doc, "{{TABLE:STATE}}")
        if state_table:
            self._fill_table(state_table, data.get('state_specs', []), ["name", "type", "initial_value", "description"])

        event_table = self._find_table_by_tag(doc, "{{TABLE:EVENT}}")
        if event_table:
            self._fill_table(event_table, data.get('event_handlers', []), ["ui_element", "trigger", "logic"])

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # --- Helper Methods ---

    def _generate_ui_structure_table(self, doc, layout_data: list, raw_components: list):
        """
        {{UI_STRUCTURE}} 위치에 [화면구성 | 유형 | 필수 | 비고] 형태의 트리 테이블 생성
        """
        target_p = self._find_and_clear_tag(doc, "{{UI_STRUCTURE}}")
        if not target_p or not layout_data: return

        # 컴포넌트 상세 정보 매핑 (Label -> Detail)
        comp_map = {c['label']: c for c in raw_components}

        # 표 생성 (헤더 + 데이터)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        target_p._p.addnext(table._tbl)

        # [헤더 설정]
        headers = ["화면 구성", "UI 유형", "필수", "비고"]
        # 열 너비 비율 (대략적): 4:2:1:3
        
        for i, text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = text
            self._set_cell_bg(cell, "E7E6E6") # 헤더 회색 배경
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # [데이터 채우기]
        for area in layout_data:
            # 1. 구역(Area) 행 추가 (Root Level)
            row = table.add_row()
            
            # Col 0: 구역명 (Bold)
            cell_area = row.cells[0]
            cell_area.text = area.get('area_name', 'Area')
            cell_area.paragraphs[0].runs[0].font.bold = True
            self._set_cell_bg(row.cells[0], "F9F9F9") # 구역 강조
            self._set_cell_bg(row.cells[1], "F9F9F9")
            self._set_cell_bg(row.cells[2], "F9F9F9")
            self._set_cell_bg(row.cells[3], "F9F9F9")

            # Col 1: Type (Area)
            row.cells[1].text = "Area"
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Col 3: 비고 (설명)
            if area.get('description'):
                row.cells[3].text = area['description']
                row.cells[3].paragraphs[0].style.font.size = Pt(9)

            # 2. 컴포넌트(Component) 행 추가 (Child Level)
            for comp_label in area.get('components', []):
                comp_detail = comp_map.get(comp_label, {})
                
                c_row = table.add_row()
                
                # Col 0: 컴포넌트명 (트리 구조 들여쓰기 적용)
                c_cell = c_row.cells[0]
                p = c_cell.paragraphs[0]
                p.text = f"   └ {comp_label}" # 공백으로 들여쓰기 시각화
                # p.paragraph_format.left_indent = Inches(0.2) # 실제 들여쓰기 (선택)
                
                # Col 1: UI 유형 (한글 매핑)
                raw_type = comp_detail.get('type', '-')
                type_map = {
                    'textbox': '입력창', 'codeview': '팝업검색', 'combo': '콤보박스',
                    'date-picker': '날짜선택', 'button': '버튼', 'grid': '그리드',
                    'textarea': '텍스트영역', 'checkbox': '체크박스', 'radio': '라디오'
                }
                ui_type = type_map.get(raw_type, raw_type)
                c_row.cells[1].text = ui_type
                c_row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Col 2: 필수 여부
                is_req = comp_detail.get('required', False)
                if is_req:
                    c_row.cells[2].text = "O"
                    c_row.cells[2].paragraphs[0].runs[0].font.bold = True
                    c_row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0) # 빨간색
                c_row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Col 3: 비고 (추가 정보가 있다면)
                # c_row.cells[3].text = ""

    def _generate_user_flow_table(self, doc, flow_data: list):
        target_p = self._find_and_clear_tag(doc, "{{USER_FLOW}}")
        if not target_p or not flow_data: return

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        target_p._p.addnext(table._tbl)

        headers = ["단계", "사용자 액션", "시스템 반응", "화면 예시"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            self._set_cell_bg(cell, "E7E6E6")
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for item in flow_data:
            row = table.add_row()
            row.cells[0].text = str(item.get('step', '-'))
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row.cells[1].text = item.get('action', '-')
            if item.get('description'):
                p = row.cells[1].add_paragraph(f"({item.get('description')})")
                p.style.font.size = Pt(8)
                p.style.font.color.rgb = RGBColor(100, 100, 100)
            
            row.cells[2].text = item.get('system_response', '-')
            row.cells[3].text = "" # 스크린샷 공간

    def _find_table_by_tag(self, doc, tag):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if tag in cell.text:
                        cell.text = cell.text.replace(tag, "")
                        return table
        return None

    def _find_and_clear_tag(self, doc, tag):
        for p in doc.paragraphs:
            if tag in p.text:
                p.text = ""
                return p
        return None

    def _set_cell_bg(self, cell, color_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color_hex)
        tcPr.append(shd)

    def _fill_table(self, table, data_list, keys):
        if not data_list: return
        for item in data_list:
            try:
                row = table.add_row().cells
                for i, key in enumerate(keys):
                    if i < len(row):
                        row[i].text = str(item.get(key, '-'))
            except: pass

    # ==========================
    # 테스트 계획서 생성
    # ==========================
    async def generate_test_plan_doc(
        self,
        screen_name: str,
        react_code: str,
        wizard_data: dict,
        images: List[dict] = None
    ) -> BytesIO:
        """테스트 계획서 생성 (Word)"""
        logger.info(f"📄 Generating test plan for: {screen_name}")
        
        # 1. LLM 데이터 추출
        try:
            prompt = get_test_plan_prompt(react_code, wizard_data)
            logger.info(f"📝 Test Plan Prompt generated: {len(prompt)} chars")
            
            text = await self._call_with_retry(prompt, "Test Plan")
            
            text = text.strip()
            if text.startswith("```json"): text = text[7:]
            if text.startswith("```"): text = text[3:]
            if text.endswith("```"): text = text[:-3]
            test_data = json.loads(text)
            logger.info("✅ LLM Test Plan Extraction Success")
        except Exception as e:
            logger.error(f"❌ LLM Test Plan Extraction Failed: {type(e).__name__}: {e}")
            traceback.print_exc()
            test_data = {"overview": {"screen_name": screen_name, "test_objective": "분석 실패"}}

        # 2. Word 생성
        return self._create_test_plan_docx(test_data, wizard_data, images)

    def _create_test_plan_docx(self, data: dict, wizard_data: dict, images: List[dict] = None) -> BytesIO:
        base_dir = os.path.dirname(os.path.dirname(__file__))
        template_path = os.path.join(base_dir, 'templates', 'test_plan_template.docx')
        
        doc = Document(template_path) if os.path.exists(template_path) else Document()

        # (A) 기본 정보 치환
        overview = data.get('overview', {})
        replacements = {
            "{{SCREEN_NAME}}": str(overview.get('screen_name', wizard_data.get('step1', {}).get('screenName', ''))),
            "{{TEST_OBJECTIVE}}": str(overview.get('test_objective', '')),
            "{{TEST_SCOPE}}": str(overview.get('test_scope', '')),
        }

        for paragraph in doc.paragraphs:
            for key, val in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(val))

        # (B) 스크린샷 삽입
        for paragraph in doc.paragraphs:
            if "{{SCREENSHOT}}" in paragraph.text:
                paragraph.text = ""
                if images:
                    for img in images:
                        try:
                            run = paragraph.add_run()
                            run.add_picture(BytesIO(img['bytes']), width=Inches(5.5))
                        except: pass
                break

        # (C) 사전조건 테이블
        precondition_table = self._find_table_by_tag(doc, "{{TABLE:PRECONDITIONS}}")
        if precondition_table:
            preconditions = overview.get('preconditions', [])
            for i, cond in enumerate(preconditions, 1):
                row = precondition_table.add_row().cells
                row[0].text = str(i)
                row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row[1].text = str(cond)

        # (D) 테스트 케이스 테이블 생성
        self._generate_test_cases_table(doc, data.get('test_cases', []))

        # (E) 테스트 시나리오 테이블
        self._generate_test_scenarios_table(doc, data.get('test_scenarios', []))

        # (F) 경계값 테스트 테이블
        boundary_table = self._find_table_by_tag(doc, "{{TABLE:BOUNDARY}}")
        if boundary_table:
            for item in data.get('boundary_tests', []):
                row = boundary_table.add_row().cells
                row[0].text = str(item.get('field', '-'))
                row[1].text = str(item.get('test_type', '-'))
                row[2].text = str(item.get('min_value', '-'))
                row[3].text = str(item.get('max_value', '-'))
                invalid = item.get('invalid_cases', [])
                row[4].text = ', '.join(invalid) if isinstance(invalid, list) else str(invalid)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _generate_test_cases_table(self, doc, test_cases: list):
        """테스트 케이스 테이블 생성"""
        target_p = self._find_and_clear_tag(doc, "{{TEST_CASES}}")
        if not target_p or not test_cases: return

        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        target_p._p.addnext(table._tbl)

        headers = ["TC ID", "분류", "테스트 항목", "테스트 절차", "예상 결과", "우선순위"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            self._set_cell_bg(cell, "E7E6E6")
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for tc in test_cases:
            row = table.add_row()
            row.cells[0].text = str(tc.get('tc_id', '-'))
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row.cells[1].text = str(tc.get('category', '-'))
            row.cells[2].text = f"{tc.get('test_item', '-')}\n{tc.get('test_description', '')}"
            
            steps = tc.get('test_steps', [])
            row.cells[3].text = '\n'.join(steps) if isinstance(steps, list) else str(steps)
            
            row.cells[4].text = str(tc.get('expected_result', '-'))
            
            priority = tc.get('priority', 'Medium')
            row.cells[5].text = str(priority)
            row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 우선순위별 색상
            if priority == 'High':
                row.cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                row.cells[5].paragraphs[0].runs[0].font.bold = True

    def _generate_test_scenarios_table(self, doc, scenarios: list):
        """테스트 시나리오 테이블 생성"""
        target_p = self._find_and_clear_tag(doc, "{{TEST_SCENARIOS}}")
        if not target_p or not scenarios: return

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        target_p._p.addnext(table._tbl)

        headers = ["시나리오 ID", "시나리오명", "설명", "테스트 절차"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            self._set_cell_bg(cell, "E7E6E6")
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for sc in scenarios:
            row = table.add_row()
            row.cells[0].text = str(sc.get('scenario_id', '-'))
            row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row.cells[1].text = str(sc.get('scenario_name', '-'))
            row.cells[2].text = str(sc.get('description', '-'))
            
            steps = sc.get('steps', [])
            if isinstance(steps, list):
                row.cells[3].text = ' → '.join(steps)
            else:
                row.cells[3].text = str(steps)

    # ==========================
    # 사용자 매뉴얼 생성
    # ==========================
    async def generate_user_manual_doc(
        self,
        screen_name: str,
        react_code: str,
        wizard_data: dict,
        images: List[dict] = None
    ) -> BytesIO:
        """사용자 매뉴얼 생성 (Word)"""
        logger.info(f"📄 Generating user manual for: {screen_name}")
        
        # 1. LLM 데이터 추출
        try:
            prompt = get_user_manual_prompt(react_code, wizard_data)
            logger.info(f"📝 User Manual Prompt generated: {len(prompt)} chars")
            
            text = await self._call_with_retry(prompt, "User Manual")
            
            text = text.strip()
            if text.startswith("```json"): text = text[7:]
            if text.startswith("```"): text = text[3:]
            if text.endswith("```"): text = text[:-3]
            manual_data = json.loads(text)
            logger.info("✅ LLM User Manual Extraction Success")
        except Exception as e:
            logger.error(f"❌ LLM User Manual Extraction Failed: {type(e).__name__}: {e}")
            traceback.print_exc()
            manual_data = {"overview": {"screen_name": screen_name, "description": "분석 실패"}}

        # 2. Word 생성
        return self._create_user_manual_docx(manual_data, wizard_data, images)

    def _create_user_manual_docx(self, data: dict, wizard_data: dict, images: List[dict] = None) -> BytesIO:
        base_dir = os.path.dirname(os.path.dirname(__file__))
        template_path = os.path.join(base_dir, 'templates', 'user_manual_template.docx')
        
        doc = Document(template_path) if os.path.exists(template_path) else Document()

        # (A) 기본 정보 치환
        overview = data.get('overview', {})
        replacements = {
            "{{SCREEN_NAME}}": str(overview.get('screen_name', wizard_data.get('step1', {}).get('screenName', ''))),
            "{{DESCRIPTION}}": str(overview.get('description', wizard_data.get('step1', {}).get('description', ''))),
            "{{TARGET_USERS}}": str(overview.get('target_users', '현업 담당자')),
        }

        for paragraph in doc.paragraphs:
            for key, val in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(val))

        # (B) 메인 스크린샷 삽입
        for paragraph in doc.paragraphs:
            if "{{SCREENSHOT_MAIN}}" in paragraph.text:
                paragraph.text = ""
                if images and len(images) > 0:
                    try:
                        run = paragraph.add_run()
                        run.add_picture(BytesIO(images[0]['bytes']), width=Inches(6.0))
                    except: pass
                break

        # (C) UI 구조 테이블 생성
        self._generate_manual_ui_table(doc, data.get('ui_structure', []))

        # (D) 수행 절차 섹션 생성
        self._generate_procedure_section(doc, data.get('procedures', []), images)

        # (E) 문제해결 테이블 생성
        self._generate_troubleshooting_table(doc, data.get('troubleshooting', []))

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _generate_manual_ui_table(self, doc, ui_structure: list):
        """사용자 매뉴얼용 UI 구조 테이블"""
        target_p = self._find_and_clear_tag(doc, "{{UI_STRUCTURE}}")
        if not target_p or not ui_structure: return

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        target_p._p.addnext(table._tbl)

        headers = ["영역", "항목", "설명", "필수"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            self._set_cell_bg(cell, "4472C4")
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for area in ui_structure:
            # 영역 행
            area_row = table.add_row()
            area_name = area.get('area_name', '')
            area_row.cells[0].text = area_name
            area_row.cells[0].paragraphs[0].runs[0].font.bold = True
            self._set_cell_bg(area_row.cells[0], "D6DCE4")
            
            area_row.cells[1].text = ""
            self._set_cell_bg(area_row.cells[1], "D6DCE4")
            
            area_row.cells[2].text = area.get('description', '')
            self._set_cell_bg(area_row.cells[2], "D6DCE4")
            
            area_row.cells[3].text = ""
            self._set_cell_bg(area_row.cells[3], "D6DCE4")

            # 컴포넌트 행
            for comp in area.get('components', []):
                comp_row = table.add_row()
                comp_row.cells[0].text = ""  # 영역 칸은 비움
                comp_row.cells[1].text = f"  • {comp.get('name', '-')}"
                comp_row.cells[2].text = comp.get('description', '')
                
                is_required = comp.get('is_required', False)
                if is_required:
                    comp_row.cells[3].text = "●"
                    comp_row.cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                comp_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _generate_procedure_section(self, doc, procedures: list, images: List[dict] = None):
        """수행 절차 섹션 생성"""
        target_p = self._find_and_clear_tag(doc, "{{PROCEDURE_SECTION}}")
        if not target_p or not procedures: return

        # 이미지 매핑 (인덱스 기반)
        image_map = {}
        if images:
            for i, img in enumerate(images):
                image_map[i] = img

        insert_point = target_p._p

        for proc in procedures:
            # 절차 제목
            title = f"{proc.get('procedure_id', '')}. {proc.get('title', '')}"
            title_p = doc.add_paragraph()
            title_run = title_p.add_run(title)
            title_run.font.size = Pt(12)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0, 51, 102)
            insert_point.addnext(title_p._p)
            insert_point = title_p._p

            # 절차 설명
            if proc.get('description'):
                desc_p = doc.add_paragraph(proc['description'])
                desc_p.style.font.size = Pt(10)
                insert_point.addnext(desc_p._p)
                insert_point = desc_p._p

            # 단계별 테이블
            steps = proc.get('steps', [])
            if steps:
                step_table = doc.add_table(rows=1, cols=3)
                step_table.style = 'Table Grid'
                insert_point.addnext(step_table._tbl)
                insert_point = step_table._tbl

                headers = ["단계", "수행 방법", "결과"]
                for i, h in enumerate(headers):
                    cell = step_table.rows[0].cells[i]
                    cell.text = h
                    self._set_cell_bg(cell, "4472C4")
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                for step in steps:
                    row = step_table.add_row()
                    row.cells[0].text = str(step.get('step', '-'))
                    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row.cells[1].text = step.get('action', '-')
                    row.cells[2].text = step.get('system_response', '-')

            # Tips
            tips = proc.get('tips', [])
            if tips:
                for tip in tips:
                    tip_p = doc.add_paragraph()
                    tip_run = tip_p.add_run(f"💡 {tip}")
                    tip_run.font.size = Pt(9)
                    tip_run.font.italic = True
                    tip_run.font.color.rgb = RGBColor(0, 102, 51)
                    insert_point.addnext(tip_p._p)
                    insert_point = tip_p._p

            # 절차 간 공백
            spacer = doc.add_paragraph()
            insert_point.addnext(spacer._p)
            insert_point = spacer._p

    def _generate_troubleshooting_table(self, doc, troubleshooting: list):
        """문제해결 테이블 생성"""
        target_p = self._find_and_clear_tag(doc, "{{TROUBLESHOOTING}}")
        if not target_p or not troubleshooting: return

        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        target_p._p.addnext(table._tbl)

        headers = ["증상", "원인", "해결 방법"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            self._set_cell_bg(cell, "C65911")
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for item in troubleshooting:
            row = table.add_row()
            row.cells[0].text = str(item.get('symptom', '-'))
            row.cells[1].text = str(item.get('cause', '-'))
            row.cells[2].text = str(item.get('solution', '-'))