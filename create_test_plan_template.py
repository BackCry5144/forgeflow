import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_test_plan_template():
    doc = Document()
    
    # 1. 스타일 설정 (한글 폰트: 맑은 고딕)
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # 2. 타이틀
    title = doc.add_heading(level=0)
    run = title.add_run('[ {{SCREEN_NAME}} ] 테스트 계획서')
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    # ---------------------------------------------------------
    # 1. 개요 (Introduction)
    # ---------------------------------------------------------
    doc.add_heading('1. 개요 및 목적', level=1)
    doc.add_paragraph('본 문서는 {{SCREEN_NAME}} 화면의 기능 검증을 위한 테스트 계획을 기술합니다.')
    
    # 요약 테이블
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    # 헤더
    cells = table.rows[0].cells
    cells[0].text = "대상 화면"
    cells[1].text = "기획 의도 및 설명"
    for cell in cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, "E7E6E6") # 회색 배경

    # 내용
    cells = table.rows[1].cells
    cells[0].text = "{{SCREEN_NAME}}"
    cells[1].text = "{{DESCRIPTION}}"
    doc.add_paragraph('')

    # ---------------------------------------------------------
    # 2. 테스트 범위 (Scope)
    # ---------------------------------------------------------
    doc.add_heading('2. 테스트 범위 (Test Scope)', level=1)
    doc.add_paragraph('■ 테스트 대상:')
    p = doc.add_paragraph('• 화면 내 UI 컴포넌트 동작 확인\n• 필수값 유효성 검사 (Validation)\n• 데이터 조회 및 CRUD 정상 동작 확인')
    p.paragraph_format.left_indent = Inches(0.2)
    
    doc.add_paragraph('■ 테스트 제외 대상:')
    p = doc.add_paragraph('• 타 시스템(SAP RFC 등) 내부 로직 검증\n• 네트워크 부하 및 성능 테스트')
    p.paragraph_format.left_indent = Inches(0.2)
    doc.add_paragraph('')

    # ---------------------------------------------------------
    # 3. 테스트 환경 (Environment)
    # ---------------------------------------------------------
    doc.add_heading('3. 테스트 환경 (Environment)', level=1)
    table_env = doc.add_table(rows=3, cols=2)
    table_env.style = 'Table Grid'
    
    data_env = [
        ("OS", "Windows 10 / 11"),
        ("Browser", "Google Chrome (Latest Version)"),
        ("해상도", "1920 x 1080 (FHD) 권장")
    ]
    
    for i, (k, v) in enumerate(data_env):
        row = table_env.rows[i]
        cell_k = row.cells[0]
        cell_k.text = k
        cell_k.paragraphs[0].runs[0].font.bold = True
        _set_cell_bg(cell_k, "F2F2F2")
        
        row.cells[1].text = v
    doc.add_paragraph('')

    # ---------------------------------------------------------
    # 4. 테스트 케이스 (Test Cases) - 핵심 🔥
    # ---------------------------------------------------------
    doc.add_heading('4. 테스트 케이스 (Test Cases)', level=1)
    doc.add_paragraph('다음은 화면의 주요 기능을 검증하기 위한 시나리오입니다.')
    
    # 🔥 코드가 이 태그를 찾아서 '테스트 케이스 테이블'을 생성해 넣습니다.
    doc.add_paragraph('{{TEST_CASE_TABLE}}')
    doc.add_paragraph('')

    # ---------------------------------------------------------
    # 5. 합격 기준 (Criteria)
    # ---------------------------------------------------------
    doc.add_heading('5. 합격 기준 (Pass/Fail Criteria)', level=1)
    doc.add_paragraph('• 정의된 모든 테스트 케이스(Happy Path)가 오류 없이 수행되어야 함.')
    doc.add_paragraph('• Critical, Major 등급의 결함이 존재하지 않아야 함.')
    doc.add_paragraph('• UI 깨짐 현상이 없어야 함.')

    # 저장
    output_path = 'backend/templates/test_plan_template.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ 테스트 계획서 템플릿 생성 완료: {output_path}")

def _set_cell_bg(cell, color_hex):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

if __name__ == "__main__":
    create_test_plan_template()