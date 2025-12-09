import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_manual_template():
    doc = Document()
    
    # 스타일 설정
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    # 타이틀
    title = doc.add_heading(level=0)
    run = title.add_run('[ {{SCREEN_NAME}} ] 사용자 매뉴얼')
    run.font.name = 'Malgun Gothic'
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 0, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    # ---------------------------------------------------------
    # 1. 화면 개요 (Overview)
    # ---------------------------------------------------------
    doc.add_heading('1. 화면 개요', level=1)
    doc.add_paragraph('본 화면은 다음 업무를 수행하기 위해 사용됩니다.')
    
    # 박스형 설명
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.rows[0].cells[0]
    cell.text = "{{DESCRIPTION}}"
    _set_cell_bg(cell, "F9F9F9")
    doc.add_paragraph('')

    # ---------------------------------------------------------
    # 2. 화면 구성 (UI Structure)
    # ---------------------------------------------------------
    doc.add_heading('2. 화면 구성', level=1)
    doc.add_paragraph('화면의 주요 구성 요소는 다음과 같습니다.')
    
    # 메인 스크린샷 자리
    p = doc.add_paragraph('{{SCREENSHOT_MAIN}}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # UI 구조 설명 (설계서에서 만든 트리 구조 재활용 가능)
    doc.add_paragraph('{{UI_STRUCTURE}}')
    doc.add_paragraph('')

    # ---------------------------------------------------------
    # 3. 주요 업무 절차 (Procedures) - 핵심 🔥
    # ---------------------------------------------------------
    doc.add_heading('3. 주요 업무 절차 (Procedures)', level=1)
    doc.add_paragraph('각 업무별 상세 수행 방법입니다.')
    doc.add_paragraph('')
    
    # 🔥 코드가 이 태그를 찾아서 'Step-by-Step 가이드'를 생성해 넣습니다.
    # (예: 3.1 조회하기 ... 3.2 저장하기 ...)
    doc.add_paragraph('{{PROCEDURE_SECTION}}')
    doc.add_paragraph('')

    # ---------------------------------------------------------
    # 4. 문제 해결 (Troubleshooting)
    # ---------------------------------------------------------
    doc.add_heading('4. 문제 해결 (Troubleshooting)', level=1)
    doc.add_paragraph('사용 중 발생할 수 있는 주요 문제와 해결 방법입니다.')
    
    table_ts = doc.add_table(rows=3, cols=2)
    table_ts.style = 'Table Grid'
    
    # 헤더
    table_ts.rows[0].cells[0].text = "현상"
    table_ts.rows[0].cells[1].text = "원인 및 조치 방법"
    for cell in table_ts.rows[0].cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, "E7E6E6")

    # 예시 데이터
    data_ts = [
        ("조회 버튼을 눌러도 반응이 없을 때", "필수 검색 조건(빨간색 별표)이 모두 입력되었는지 확인하세요."),
        ("엑셀 다운로드 파일이 열리지 않을 때", "파일 확장자가 .xlsx인지 확인하고, 엑셀 프로그램 버전 호환성을 확인하세요.")
    ]
    
    for i, (k, v) in enumerate(data_ts):
        row = table_ts.rows[i+1]
        row.cells[0].text = k
        row.cells[1].text = v

    # 저장
    output_path = 'backend/templates/user_manual_template.docx'
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"✅ 사용자 매뉴얼 템플릿 생성 완료: {output_path}")

def _set_cell_bg(cell, color_hex):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

if __name__ == "__main__":
    create_manual_template()